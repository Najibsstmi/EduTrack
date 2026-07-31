from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parent
REFERENCE_DOCX = BASE_DIR / "reference-minit-curai-pbd-bahasa-melayu-2026.docx"
OUTPUT_DOCX = BASE_DIR / "MINIT CURAI PBD PPT PANITIA SAINS 2026.docx"


SCIENCE_DATA = {
    "overall": {"total": 320, "tp3": 149, "tp3p": 46.6, "tp4": 152, "tp4p": 47.5, "tp5": 19, "tp5p": 5.9},
    "grades": {
        1: {"total": 83, "tp3": 70, "tp3p": 84.3, "tp4": 10, "tp4p": 12.0, "tp5": 3, "tp5p": 3.6},
        2: {"total": 72, "tp3": 4, "tp3p": 5.6, "tp4": 65, "tp4p": 90.3, "tp5": 3, "tp5p": 4.2},
        3: {"total": 64, "tp3": 45, "tp3p": 70.3, "tp4": 16, "tp4p": 25.0, "tp5": 3, "tp5p": 4.7},
        4: {"total": 52, "tp3": 11, "tp3p": 21.2, "tp4": 35, "tp4p": 67.3, "tp5": 6, "tp5p": 11.5},
        5: {"total": 49, "tp3": 19, "tp3p": 38.8, "tp4": 26, "tp4p": 53.1, "tp5": 4, "tp5p": 8.2},
    },
}


ATTENDANCE = [
    "Nuremelia binti Mohd Rawan (Guru Sains: 4 SONGKET, 5 BUDAYAWAN, 2 MASRI)",
    "Mohd Najib bin Jaafar (Guru Sains: 4 CINDAI, 5 SENIMAN, 2 MASRI, 2 INANG)",
    "Mazlita binti Md Salleh (Guru Sains: 3 BANGSAWAN, 4 SUTERA, 5 KARYAWAN, 3 BONEKA)",
    "Nurjannah binti Berhanuddin (Guru Sains: 1 TANGO, 1 JAZZ, 1 BALADA, 3 MUZIKAL)",
]


def clear_cell(cell):
    cell._tc.clear_content()


def add_paragraph(cell, text="", bold=False, font_size=10):
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    return paragraph


def set_cell_lines(cell, lines, font_size=10):
    clear_cell(cell)
    for item in lines:
        if isinstance(item, tuple):
            text, bold = item
        else:
            text, bold = item, False
        add_paragraph(cell, text, bold=bold, font_size=font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_cell_text(cell, text, bold=False, font_size=10):
    set_cell_lines(cell, [(text, bold)], font_size=font_size)


def pct(value):
    return f"{value:.1f}%"


def distribution_sentence(total, tp3, tp3p, tp4, tp4p, tp5, tp5p):
    parts = [
        f"Daripada {total} orang murid, seramai {tp3} orang ({pct(tp3p)}) berada pada TP3",
        f"{tp4} orang ({pct(tp4p)}) pada TP4",
        f"{tp5} orang ({pct(tp5p)}) mencapai TP5",
    ]
    return ", ".join(parts) + ". Tiada murid berada pada TP1, TP2, TP6 atau Tidak Ditaksir (TD)."


def overall_lines():
    data = SCIENCE_DATA["overall"]
    return [
        "Merujuk analisis pencapaian PBD (PPT) murid tahun 2026",
        ("Sains", True),
        (
            "Secara keseluruhannya, seramai 320 orang murid telah ditaksir. "
            "Majoriti murid iaitu 152 orang (47.5%) berada pada TP4, diikuti "
            "149 orang (46.6%) pada TP3 dan 19 orang (5.9%) pada TP5. "
            "Tiada murid berada pada TP1, TP2, TP6 atau Tidak Ditaksir (TD)."
        ),
        ("Analisis Keseluruhan:", True),
        (
            "Dapatan ini menunjukkan tahap penguasaan Sains adalah baik kerana semua murid "
            "telah mencapai tahap minimum TP3 hingga TP6. Walau bagaimanapun, hampir separuh "
            "murid masih berada pada TP3, maka penekanan perlu diberikan kepada penguasaan "
            "konsep asas, kemahiran proses sains, penaakulan dan aplikasi konsep supaya murid "
            "dapat bergerak ke TP4 dan TP5. Bilangan TP5 masih kecil dan boleh dipertingkatkan "
            "melalui aktiviti pengayaan serta tugasan penyiasatan yang lebih mencabar."
        ),
        ("Cadangan Intervensi Keseluruhan:", True),
        "Melaksanakan intervensi berfokus bagi murid TP3 melalui latihan konsep asas, istilah sains dan kemahiran proses sains.",
        "Mengadakan aktiviti eksperimen ringkas, simulasi, pembelajaran berasaskan inkuiri dan projek mini bagi mengukuhkan aplikasi konsep.",
        "Memberikan pengayaan KBAT dan tugasan penyiasatan kepada murid TP4 untuk melonjak ke TP5.",
        "Memantau perkembangan PBD secara berkala melalui perbincangan profesional panitia dan semakan eviden murid.",
    ]


def grade_lines(grade):
    d = SCIENCE_DATA["grades"][grade]
    if grade == 1:
        return [
            ("Tingkatan 1", True),
            distribution_sentence(d["total"], d["tp3"], d["tp3p"], d["tp4"], d["tp4p"], d["tp5"], d["tp5p"]),
            ("Analisis:", True),
            (
                "Pencapaian Tingkatan 1 menunjukkan majoriti murid masih berada pada TP3. "
                "Kelas 1 BALADA, 1 JAZZ dan 1 TANGO masing-masing didominasi TP3, menandakan "
                "murid telah menguasai asas Sains tetapi masih memerlukan bimbingan untuk "
                "mengaplikasikan pengetahuan dalam situasi baharu."
            ),
            ("Kelemahan Calon:", True),
            "Penguasaan istilah Sains, pemerhatian dan inferens masih perlu diperkukuh.",
            "Sebahagian murid masih kurang yakin menjelaskan hubungan sebab dan akibat secara saintifik.",
            "Kemahiran merekod pemerhatian dan mentafsir maklumat daripada rajah atau jadual masih sederhana.",
            ("Intervensi:", True),
            "Memperbanyak aktiviti amali ringkas, kuiz konsep dan latihan kosa kata Sains.",
            "Memberikan lembaran berperingkat mengikut tahap penguasaan murid.",
            "Menggunakan pembelajaran koperatif dan bimbingan rakan sebaya untuk meningkatkan keyakinan murid.",
        ]
    if grade == 2:
        return [
            ("Tingkatan 2", True),
            distribution_sentence(d["total"], d["tp3"], d["tp3p"], d["tp4"], d["tp4p"], d["tp5"], d["tp5p"]),
            ("Analisis:", True),
            (
                "Majoriti murid Tingkatan 2 berada pada TP4, khususnya kelas 2 INANG dan 2 MASRI "
                "yang menunjukkan penguasaan konsep yang kukuh. Namun, peratus TP5 masih rendah "
                "dan masih terdapat sebahagian kecil murid pada TP3 yang memerlukan pemulihan berfokus."
            ),
            ("Kelemahan Calon:", True),
            "Sebahagian murid masih kurang mantap dalam menghuraikan jawapan berdasarkan bukti.",
            "Murid belum konsisten mentafsir graf, jadual dan pemboleh ubah eksperimen.",
            "Pengayaan KBAT masih perlu diperkukuh untuk melonjakkan murid TP4 ke TP5.",
            ("Intervensi:", True),
            "Menjalankan bimbingan berfokus kepada murid TP3 melalui latihan topikal.",
            "Memperbanyak latihan tafsiran data, graf dan soalan aplikasi harian.",
            "Memberikan tugasan inkuiri dan soalan KBAT kepada murid TP4 yang berpotensi mencapai TP5.",
        ]
    if grade == 3:
        return [
            ("Tingkatan 3", True),
            distribution_sentence(d["total"], d["tp3"], d["tp3p"], d["tp4"], d["tp4p"], d["tp5"], d["tp5p"]),
            ("Analisis:", True),
            (
                "Prestasi Tingkatan 3 masih perlu diperkukuh kerana 45 orang murid berada pada TP3. "
                "Kelas 3 MUZIKAL keseluruhannya berada pada TP3, manakala 3 BANGSAWAN menunjukkan "
                "taburan yang lebih baik dengan kewujudan murid TP5. Fokus perlu diberikan kepada "
                "aplikasi konsep dan kemahiran proses sains sebelum murid memasuki tahap menengah atas."
            ),
            ("Kelemahan Calon:", True),
            "Murid kurang menguasai kemahiran mengenal pasti pemboleh ubah dan membuat inferens.",
            "Sebahagian murid sukar mengaitkan konsep Sains dengan fenomena atau situasi harian.",
            "Jawapan masih kurang lengkap dari segi huraian, bukti dan penggunaan istilah tepat.",
            ("Intervensi:", True),
            "Melaksanakan modul pemulihan mengikut tema dan topik yang sukar dikuasai.",
            "Memperbanyak aktiviti laporan amali, perbincangan dapatan dan refleksi eksperimen.",
            "Mengadakan latihan KBAT secara berkala untuk meningkatkan keupayaan aplikasi konsep.",
        ]
    if grade == 4:
        return [
            ("Tingkatan 4", True),
            distribution_sentence(d["total"], d["tp3"], d["tp3p"], d["tp4"], d["tp4p"], d["tp5"], d["tp5p"]),
            ("Analisis:", True),
            (
                "Pencapaian Tingkatan 4 adalah baik dengan majoriti murid berada pada TP4. "
                "Kelas 4 CINDAI menonjol dengan 6 orang murid TP5, manakala 4 SONGKET hampir "
                "keseluruhannya berada pada TP4. Kelas 4 SUTERA memerlukan perhatian kerana "
                "sebahagian besar murid masih pada TP3."
            ),
            ("Kelemahan Calon:", True),
            "Murid masih memerlukan pengukuhan bagi konsep abstrak dan penaakulan saintifik.",
            "Perancangan penyiasatan, tafsiran data dan penulisan kesimpulan masih belum konsisten.",
            "Sebahagian murid bergantung kepada contoh guru apabila menjawab soalan aplikasi.",
            ("Intervensi:", True),
            "Mengadakan latihan berfokus kepada kemahiran proses sains dan soalan berformat SPM.",
            "Melaksanakan bimbingan kumpulan kecil bagi murid TP3, khususnya mengikut topik lemah.",
            "Memberikan tugasan pengayaan eksperimen dan kajian kes kepada murid TP4 dan TP5.",
        ]
    return [
        ("Tingkatan 5", True),
        distribution_sentence(d["total"], d["tp3"], d["tp3p"], d["tp4"], d["tp4p"], d["tp5"], d["tp5p"]),
        ("Analisis:", True),
        (
            "Prestasi Tingkatan 5 menunjukkan lebih separuh murid berada pada TP4 dan 4 orang murid "
            "mencapai TP5. Kelas 5 BUDAYAWAN memperlihatkan pencapaian paling kukuh dengan 4 orang "
            "TP5, manakala 5 KARYAWAN dan 5 SENIMAN masih memerlukan bimbingan untuk mengurangkan "
            "bilangan murid TP3 menjelang pentaksiran seterusnya."
        ),
        ("Kelemahan Calon:", True),
        "Murid masih kurang konsisten mengaplikasikan konsep Sains dalam soalan stimulus dan situasi baharu.",
        "Kemahiran menghuraikan jawapan dengan fakta, sebab dan contoh masih perlu diperkukuh.",
        "Sebahagian murid belum menguasai strategi ulang kaji berfokus dan analisis kesilapan.",
        ("Intervensi:", True),
        "Mengadakan kelas pecutan akhir, latih tubi topikal dan perbincangan soalan berformat SPM.",
        "Melaksanakan kuiz diagnostik serta analisis kesilapan bagi menentukan topik intervensi.",
        "Menjalankan sesi mentor-mentee dan bimbingan individu bagi murid TP3 serta pengayaan bagi murid TP4.",
    ]


def clone_paragraph_style(source_cell, target_cell):
    if not source_cell.paragraphs or not target_cell.paragraphs:
        return
    source_p = source_cell.paragraphs[0]
    for paragraph in target_cell.paragraphs:
        paragraph.alignment = source_p.alignment
        paragraph.paragraph_format.left_indent = source_p.paragraph_format.left_indent
        paragraph.paragraph_format.right_indent = source_p.paragraph_format.right_indent
        paragraph.paragraph_format.first_line_indent = source_p.paragraph_format.first_line_indent


def remove_template_tail_sections(document):
    body = document._body._element
    children = list(body)
    table_index = next(index for index, child in enumerate(children) if child.tag == qn("w:tbl"))

    preferred_section = None
    for child in children[table_index + 1 :]:
        if child.tag == qn("w:p"):
            paragraph_props = child.find(qn("w:pPr"))
            if paragraph_props is not None:
                section_props = paragraph_props.find(qn("w:sectPr"))
                if section_props is not None:
                    preferred_section = deepcopy(section_props)
                    break

    if preferred_section is None:
        existing_section = body.find(qn("w:sectPr"))
        preferred_section = deepcopy(existing_section) if existing_section is not None else None

    for child in list(body)[table_index + 1 :]:
        body.remove(child)

    if preferred_section is not None:
        body.append(preferred_section)


def main():
    document = Document(str(REFERENCE_DOCX))
    table = document.tables[0]

    # Top metadata.
    set_cell_text(table.cell(0, 2), "Penyelarasan dan Jaminan Kualiti Markah PBD Pertengahan Tahun peringkat Panitia Sains")
    set_cell_text(table.cell(0, 4), "17 Julai 2026")
    set_cell_text(table.cell(1, 4), "10.30 pagi - 11.30 pagi")
    set_cell_text(table.cell(2, 2), "Bilik Peperiksaan, Sekolah Seni Malaysia Johor.")

    attendance_lines = ["KEHADIRAN :"] + [f"{index}. {name}" for index, name in enumerate(ATTENDANCE, start=1)]
    set_cell_lines(table.cell(3, 0), attendance_lines)

    # Issue headers and content.
    set_cell_text(table.cell(5, 1), "Penyelarasan Markah PBD Sains Keseluruhan", bold=True)
    set_cell_lines(table.cell(6, 1), overall_lines())
    set_cell_text(table.cell(6, 4), "Maklum")

    set_cell_text(table.cell(7, 1), "Analisis PBD Sains Mengikut Tahap Penguasaan (TP) dalam Tingkatan Tahun 2026", bold=True)
    for row_index, grade in zip(range(8, 13), range(1, 6)):
        set_cell_lines(table.cell(row_index, 1), grade_lines(grade))

    set_cell_lines(table.cell(13, 0), [
        "Nama Pencatat :",
        "",
        "",
        "....................................",
        "Nuremelia binti Mohd Rawan",
    ])
    set_cell_lines(table.cell(13, 4), ["Tarikh Disediakan:", "", "", "17 Julai 2026"])

    # Keep typography tidy after adding fresh paragraphs.
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    if run.font.size is None:
                        run.font.size = Pt(10)

    remove_template_tail_sections(document)
    document.save(str(OUTPUT_DOCX))
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
